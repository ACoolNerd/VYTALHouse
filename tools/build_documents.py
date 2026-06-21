#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"

GOLD = RGBColor(200, 164, 107)
DARK = RGBColor(17, 23, 34)
INK = RGBColor(34, 38, 46)
MUTED = RGBColor(92, 98, 110)


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_font(run, size: int | float = 11, color: RGBColor = INK, bold: bool = False) -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, 17 if level == 1 else 13, DARK if level == 1 else INK, True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(text), 10.5, INK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(item), 10.5, INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_fill(cell, "111722")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(header), 9, RGBColor(255, 255, 255), True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(value)), 9, INK)
    doc.add_paragraph()


def build_dossier() -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("VYTAL HOUSE"), 26, DARK, True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run("Recharge. Recover. Evolve."), 13, GOLD, True)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(meta.add_run("Full Manifest Dossier | Local Draft Package | June 19, 2026"), 9.5, MUTED)

    add_heading(doc, "Executive Control Summary")
    add_body(
        doc,
        "VYTAL House is structured as an isolated recovery and wellness club package with source-backed planning, governed data records, local-first artifact outputs, and explicit gates for legal, medical, financial, vendor, payment, Notion, and Figma actions.",
    )
    add_bullets(
        doc,
        [
            "Brand is locked to VYTAL House, the tagline Recharge. Recover. Evolve., and owners Chauncey Gardner and Kathy Ha.",
            "The software MVP uses Next.js, React, Firebase-ready rules, seeded records, and safe prototype forms.",
            "The operating package includes Notion importables, Figma tokens, vendor outreach, social calendar, source manifest, and delivery manifest.",
            "The quality gate requires a score of at least 97/100 before runtime handoff.",
        ],
    )

    add_heading(doc, "Membership Baseline")
    add_table(
        doc,
        ["Tier", "Monthly Price", "Projected Members", "Projected MRR", "Role"],
        [
            ["Core", "$299", "150", "$44,850", "Entry recovery access"],
            ["Elite", "$599", "100", "$59,900", "Unlimited recovery and IV/NAD+ credit"],
            ["Black", "$1,499", "50", "$74,950", "Concierge wellness and diagnostics"],
        ],
    )

    add_heading(doc, "Service System")
    add_table(
        doc,
        ["Service", "Purpose", "Gate"],
        [
            ["Hyperbaric Oxygen Therapy", "Anchor clinical recovery modality", "Clinical review"],
            ["IV and NAD+ Lounge", "High-touch lounge treatment flow", "Medical director review"],
            ["Red Light Therapy", "Photobiomodulation and recovery", "Protocol review"],
            ["Cryotherapy", "Cold exposure recovery lane", "Safety review"],
            ["Sauna and Cold Plunge", "Thermal contrast suite", "Operations review"],
            ["Float and Dry Float", "Restoration and stress relief", "Operations review"],
        ],
    )

    add_heading(doc, "Separation of Duties")
    add_table(
        doc,
        ["Role", "Can Own", "Cannot Do Without Approval"],
        [
            ["owner", "Final launch gates", "Bypass professional review"],
            ["admin", "Platform and quality evidence", "Collect protected health data"],
            ["clinical", "Medical protocols", "Publish legal or financial claims"],
            ["operations", "Vendors and booking readiness", "Send vendor messages without approval"],
            ["marketing", "Content and social calendar", "Make clinical claims"],
            ["vendor", "Own quote responses", "Access member data"],
            ["member", "Own booking interest", "Access admin surfaces"],
        ],
    )

    add_heading(doc, "7-Field Schema")
    add_body(doc, "Every database-like object must include id, entity, type, name, status, owner, and updatedAt. Domain payloads live under metadata.")

    add_heading(doc, "Source and Compliance Notes")
    add_bullets(
        doc,
        [
            "Vendor and modality claims are anchored to source URLs listed in docs/01_SOURCE_MANIFEST.md.",
            "Maryland entity formation details require SDAT/Business Express review before filing.",
            "HHS HIPAA guidance is treated as a privacy reference; no PHI is collected in the prototype.",
            "The CAD package is conceptual only and requires architect/engineer conversion before permits.",
        ],
    )

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("VYTAL House | Draft package for professional review"), 8, MUTED)

    out = OUT / "VYTAL_House_Full_Manifest_Dossier.docx"
    doc.save(out)
    return out


def new_branded_doc(title_text: str, subtitle_text: str, meta_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run(title_text), 23, DARK, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run(subtitle_text), 12, GOLD, True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(meta.add_run(meta_text), 9.5, MUTED)
    return doc


def add_footer(doc: Document, text: str) -> None:
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run(text), 8, MUTED)


def build_business_plan() -> Path:
    doc = new_branded_doc(
        "VYTAL HOUSE BUSINESS PLAN",
        "Recharge. Recover. Evolve.",
        "Draft operating plan | June 19, 2026 | Professional review required",
    )

    add_heading(doc, "Executive Summary")
    add_body(
        doc,
        "VYTAL House is a premium recovery and wellness club for high-performance members in Maryland. The concept combines clinical-grade recovery modalities, a recurring membership model, concierge operations, and a dark cinematic lounge experience.",
    )
    add_bullets(
        doc,
        [
            "Brand: VYTAL House only.",
            "Owners: Chauncey Gardner and Kathy Ha.",
            "Promise: Recharge. Recover. Evolve.",
            "Guardrails: no live payments, no protected health information, no vendor sending, and no legal filing without explicit approval.",
        ],
    )

    add_heading(doc, "Market Positioning")
    add_table(
        doc,
        ["Dimension", "Plan"],
        [
            ["Category", "Recovery, longevity, and wellness club"],
            ["Primary users", "Founders, artists, athletes, executives, wellness clients, and recovery-focused members"],
            ["Differentiator", "Design-forward facility, recurring membership revenue, clinical oversight, and high-touch concierge flow"],
            ["Experience", "Private, cinematic, quiet, and operationally controlled"],
        ],
    )

    add_heading(doc, "Services")
    add_table(
        doc,
        ["Service", "Commercial Role", "Control Gate"],
        [
            ["Hyperbaric Oxygen Therapy", "Anchor clinical recovery modality", "Clinical protocol review"],
            ["IV and NAD+ Lounge", "Premium treatment lane", "Medical director and pharmacy review"],
            ["Red Light Therapy", "Photobiomodulation and recovery", "Protocol and claims review"],
            ["Cryotherapy", "Cold exposure lane", "Safety and waiver review"],
            ["Sauna and Cold Plunge", "Thermal contrast suite", "Operations and maintenance review"],
            ["Float and Dry Float", "Restoration and stress relief", "Operations review"],
            ["Compression and Recovery Zone", "Member utilization and retention", "Staff training review"],
        ],
    )

    doc.add_page_break()
    add_heading(doc, "Membership Model")
    add_table(
        doc,
        ["Tier", "Price", "Projected Members", "Projected MRR", "Includes"],
        [
            ["Core", "$299/mo", "150", "$44,850", "Red light, compression, recovery lounge, education"],
            ["Elite", "$599/mo", "100", "$59,900", "Unlimited recovery services, monthly IV/NAD+ credit, booking priority"],
            ["Black", "$1,499/mo", "50", "$74,950", "Concierge scheduling, quarterly labs, peptide optimization, VIP access"],
        ],
    )

    add_heading(doc, "Financial Baseline")
    add_bullets(
        doc,
        [
            "Target membership revenue: $179,700/month at the modeled tier mix.",
            "Additional revenue lanes: IV/NAD+, peptides and hormone programs, float sessions, cryotherapy, sauna/cold plunge, compression, events, and retail.",
            "Startup cost planning range: $750,000 to $1,800,000 depending on lease terms, tenant improvements, equipment mix, staffing, launch marketing, and working capital.",
            "All financial assumptions require accountant, tax, lending, and securities review before investor or lender use.",
        ],
    )

    add_heading(doc, "Execution Roadmap")
    add_table(
        doc,
        ["Phase", "Priority", "Owner Role", "Gate"],
        [
            ["Lease and due diligence", "Confirm site, zoning, utilities, and tenant improvement terms", "owner", "Legal and lease review"],
            ["Medical structure", "Define clinical roles, protocols, consent, and pharmacy workflows", "clinical", "Medical counsel approval"],
            ["Vendor procurement", "Request quotes, install needs, warranties, and maintenance terms", "operations", "Owner quote approval"],
            ["Digital platform", "Launch website, member app, admin, vendor CRM, and source-backed content", "admin", "Quality score >= 97"],
            ["Founder launch", "Prepare presale material and social calendar", "marketing", "Claims and payment approval"],
            ["Controlled opening", "Open with a limited member cohort and track utilization", "operations", "Staff and compliance readiness"],
        ],
    )

    add_heading(doc, "Source Appendix")
    add_bullets(
        doc,
        [
            "Oxycell, Hyperbaric Pro, Oxynergy, TrueREST, VacuActiv, Eleve, and Awesome Design MD are used as reference inputs, not endorsements.",
            "Maryland Business Express and Maryland SDAT sources are used for formation workflow planning.",
            "HHS HIPAA guidance is used as a privacy planning reference. The prototype does not collect PHI.",
        ],
    )

    add_footer(doc, "VYTAL House Business Plan | Draft for professional review")
    out = OUT / "VYTAL_House_Business_Plan.docx"
    doc.save(out)
    return out


def build_brd() -> Path:
    doc = new_branded_doc(
        "VYTAL HOUSE BRD",
        "Business Requirements Document",
        "Platform scope | June 19, 2026 | Draft for stakeholder review",
    )

    add_heading(doc, "Goal")
    add_body(
        doc,
        "Launch a full-stack VYTAL House platform that supports marketing, lead capture, booking interest, vendor procurement, internal operations, source-backed planning, and quality-gated launch control.",
    )

    add_heading(doc, "In Scope")
    add_bullets(
        doc,
        [
            "Public website for brand, services, memberships, facility concept, and contact.",
            "Member app surface for seeded bookings, session requests, education, and membership context.",
            "Admin dashboard for launch readiness, quality scoring, leads, bookings, vendors, content, and documents.",
            "Vendor CRM with procurement status, contact records, quote stage, and outreach template storage.",
            "Notion-ready Markdown/CSV importables and Figma-ready tokens/assets.",
            "Firebase-ready security rules, indexes, storage rules, and governed seed data.",
        ],
    )

    add_heading(doc, "Out of Scope Until Approval")
    add_bullets(
        doc,
        [
            "Live payment processing.",
            "Live vendor email sending.",
            "Protected health information collection.",
            "Legal filings or state portal submissions.",
            "Live Notion or Figma workspace creation.",
        ],
    )

    add_heading(doc, "Governed Data Model")
    add_body(doc, "Every database-like record uses the required fields id, entity, type, name, status, owner, and updatedAt. Domain payloads are stored in metadata.")
    add_table(
        doc,
        ["Collection", "Purpose", "Primary Role", "Required Gate"],
        [
            ["members", "Member profile and membership state", "member", "No PHI without approval"],
            ["bookings", "Booking interest and operational scheduling", "operations", "Audit event required"],
            ["services", "Source-backed service catalog", "clinical", "Claims review"],
            ["vendors", "Vendor pipeline and procurement", "operations", "No sending without approval"],
            ["contentItems", "Website and wiki content", "marketing", "Brand isolation review"],
            ["socialPosts", "Social calendar", "marketing", "Claims review"],
            ["leads", "Interest capture", "admin", "No sensitive medical data"],
            ["documents", "Controlled artifact register", "admin", "Professional review tags"],
            ["auditEvents", "Governance event trail", "admin", "Append-only pattern"],
            ["qualityScores", "Pre-run quality gate", "owner", "Minimum 97/100"],
        ],
    )

    doc.add_page_break()
    add_heading(doc, "Role and Route Matrix")
    add_table(
        doc,
        ["Surface", "Allowed Roles", "Controls"],
        [
            ["Public website", "owner, admin, marketing, operations, member", "Read-only content and lead form"],
            ["Member app", "owner, admin, member", "Prototype booking interest only"],
            ["Admin dashboard", "owner, admin, operations", "Quality and launch readiness only"],
            ["Clinical review", "owner, clinical", "Protocol and claims governance"],
            ["Vendor CRM", "owner, admin, operations, vendor", "No outbound email mutation"],
            ["Content calendar", "owner, admin, marketing", "No unreviewed clinical claims"],
        ],
    )

    add_heading(doc, "Functional Requirements")
    add_bullets(
        doc,
        [
            "Users can browse services, memberships, facility zones, and source-backed reference content.",
            "Leads can submit interest through a safe endpoint with no medical intake fields.",
            "Members can view seeded bookings and request booking interest.",
            "Admins can scan quality evidence and operational readiness records.",
            "Operations can manage vendor quote stages and source notes.",
            "Marketing can manage social posts and content readiness.",
        ],
    )

    add_heading(doc, "Non-Functional Requirements")
    add_bullets(
        doc,
        [
            "VYTAL-only brand isolation.",
            "Responsive desktop and mobile UI.",
            "Accessible contrast, keyboard-friendly controls, and no known text overlap.",
            "No unreviewed clinical, legal, or financial claims.",
            "Build, typecheck, render, and quality gate evidence must be stored locally.",
        ],
    )

    add_footer(doc, "VYTAL House BRD | Draft for stakeholder and professional review")
    out = OUT / "VYTAL_House_Business_Requirements_Document.docx"
    doc.save(out)
    return out


def build_articles_draft() -> Path:
    doc = new_branded_doc(
        "VYTAL HOUSE ENTITY FORMATION STACK",
        "Articles of Organization Templates",
        "Maryland filing drafts | June 19, 2026 | Not legal advice",
    )

    add_heading(doc, "Legal Notice")
    add_body(
        doc,
        "This document contains planning drafts only. It is not legal advice and must be reviewed by qualified counsel before filing with Maryland Department of Assessments and Taxation (SDAT) or through Maryland Business Express.",
    )

    add_heading(doc, "Filing Data To Confirm")
    add_table(
        doc,
        ["Field", "Status", "Control"],
        [
            ["Entity names and availability", "To confirm", "Owner and counsel review"],
            ["Principal Maryland office addresses", "To confirm", "Do not publish until approved"],
            ["Resident agent name and physical Maryland address", "To confirm", "Resident agent consent required"],
            ["Organizer signatures", "To confirm", "Owner approval required"],
            ["Resident agent signatures", "To confirm", "Agent approval required"],
            ["Return addresses and filing fees", "To confirm", "SDAT/Business Express verification"],
        ],
    )

    entities = [
        "VYTAL Holdings LLC",
        "VYTAL Operations LLC",
        "VYTAL Intellectual Property LLC",
        "VYTAL Property Holdings LLC",
        "VYTAL House LLC"
    ]

    for index, entity_name in enumerate(entities):
        if index > 0:
            doc.add_page_break()
        
        add_heading(doc, f"Articles of Organization for {entity_name}")
        
        articles = [
            ("Article I - Name", f"The name of the limited liability company is {entity_name}."),
            ("Article II - Formation and Legal Authority", "The company is organized under the laws of the State of Maryland and any applicable Maryland limited liability company provisions."),
            ("Article III - Purpose", f"The company may engage in any lawful business activity permitted under Maryland law, serving its role in the VYTAL House platform architecture ({'holding equity in subsidiaries' if 'Holdings' in entity_name else 'managing spa services, booking, and staff' if 'Operations' in entity_name else 'holding trademarks, software copyright, and IP' if 'Intellectual Property' in entity_name else 'holding leasehold interests and physical improvements' if 'Property' in entity_name else 'membership-based recovery and wellness brand operations'}). Any clinical service model requires professional review and compliant operating controls."),
            ("Article IV - Principal Office", "The principal office shall be located at a Maryland address to be confirmed before filing. Public drafts should not include private or unconfirmed addresses."),
            ("Article V - Resident Agent", "The resident agent and resident office shall be confirmed before filing. The resident agent must consent to service of process responsibilities."),
            ("Article VI - Duration", "The company shall continue perpetually unless dissolved according to law and any adopted operating agreement."),
            ("Article VII - Management", "The company shall be managed by its members or managers as defined in an operating agreement adopted by the members."),
            ("Article VIII - Membership", "Membership interests, admission of additional members, capital contributions, voting rights, and transfer restrictions shall be controlled by the operating agreement."),
            ("Article IX - Indemnification", "The company may indemnify members, managers, officers, and agents to the fullest extent permitted by Maryland law, subject to professional legal review."),
            ("Article X - Effective Date", "The articles become effective upon acceptance for filing by the Maryland Department of Assessments and Taxation unless a permitted delayed effective date is specified."),
        ]
        for heading, body in articles:
            add_heading(doc, heading, level=2)
            add_body(doc, body)

    doc.add_page_break()
    add_heading(doc, "Draft Controls")
    add_bullets(
        doc,
        [
            "Do not file these drafts without explicit owner approval.",
            "Do not treat these drafts as counsel-approved legal language.",
            "Keep signed final documents outside public repository paths unless counsel approves storage controls.",
            "Maintain separate operating agreements, consent forms, insurance evidence, and clinical governance records before launch.",
        ],
    )

    add_footer(doc, "VYTAL House Articles Stack Drafts | Not legal advice")
    out = OUT / "VYTAL_House_Articles_of_Organization_Draft.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    built = [
        build_dossier(),
        build_business_plan(),
        build_brd(),
        build_articles_draft(),
    ]
    for path in built:
        print(path)
